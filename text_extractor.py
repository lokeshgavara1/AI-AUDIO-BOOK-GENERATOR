"""
Text Extractor Module
Handles text extraction from PDF, DOCX, and TXT files
"""

import PyPDF2
import pdfplumber
from docx import Document
import streamlit as st
from typing import Optional


class TextExtractor:
    """Extract text from various file formats"""
    
    @staticmethod
    def extract_from_pdf(file) -> str:
        """
        Extract text from PDF file using both PyPDF2 and pdfplumber
        
        Args:
            file: Uploaded file object
            
        Returns:
            str: Extracted text
        """
        try:
            text = ""
            
            # Try with pdfplumber first (better for complex PDFs)
            try:
                with pdfplumber.open(file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                st.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
                
                # Fallback to PyPDF2
                file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
        
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file: Uploaded file object
            
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(file)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_from_txt(file) -> str:
        """
        Extract text from TXT file
        
        Args:
            file: Uploaded file object
            
        Returns:
            str: Extracted text
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    text = file.read().decode(encoding)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with supported encodings")
        
        except Exception as e:
            raise Exception(f"Error extracting text from TXT: {str(e)}")
    
    @staticmethod
    def extract_text(file, file_type: str) -> str:
        """
        Main extraction method that routes to appropriate extractor
        
        Args:
            file: Uploaded file object
            file_type: File extension (pdf, docx, txt)
            
        Returns:
            str: Extracted text
        """
        file_type = file_type.lower()
        
        if file_type == "pdf":
            return TextExtractor.extract_from_pdf(file)
        elif file_type == "docx":
            return TextExtractor.extract_from_docx(file)
        elif file_type == "txt":
            return TextExtractor.extract_from_txt(file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def chunk_text(text: str, max_chunk_size: int = 4000) -> list:
        """
        Split long text into chunks for processing
        
        Args:
            text: Input text
            max_chunk_size: Maximum characters per chunk
            
        Returns:
            list: List of text chunks
        """
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        sentences = text.replace('\n', ' ').split('. ')
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 2 <= max_chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
